#!/usr/bin/env python3
"""Generate a Word report from GPS data.

Steps:
 1. Run chart.py to produce SVG charts (speed and deceleration charts).
 2. Run passingHome.py to produce the halt arrival table (markdown style on stdout).
 3. Parse the table and embed it plus the charts into a Word (.docx) document.

Usage:
  python generate_report.py path/to/data.csv --output report.docx [--loco L123] [--min-stop 30]

Notes:
 - Requires python-docx. If missing the script will attempt to install it (pip).
 - SVG is not directly supported by python-docx, so an attempt is made to convert each SVG to PNG using cairosvg.
   If cairosvg is not available, those images will be skipped with a warning.
"""
import sys, os, subprocess, glob, io, time, shutil, tempfile, re
from typing import List, Tuple, Optional
import datetime

# ---------------------------------------------------------------------------

def ensure_package(module: str, pip_name: Optional[str] = None):
    """Ensure a module is importable; if not, install via pip_name (or module if pip_name omitted)."""
    try:
        __import__(module)
        return True
    except ImportError:
        pkg = pip_name or module
        print(f'Installing missing package: {pkg} ...')
        r = subprocess.run([sys.executable, '-m', 'pip', 'install', pkg], stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
        if r.returncode != 0:
            print(f'Failed to install {pkg}:\n{r.stdout}')
            return False
        try:
            __import__(module)
            return True
        except ImportError:
            return False

# ---------------------------------------------------------------------------

def run_script_collect_output(script: str, args: List[str]) -> Tuple[int, str]:
    cmd = [sys.executable, script] + args
    print('Running:', ' '.join(cmd))
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
    if proc.returncode != 0:
        print(f'Warning: {script} exited with code {proc.returncode}')
    return proc.returncode, proc.stdout

# ---------------------------------------------------------------------------

def parse_markdown_table(md: str) -> Tuple[List[str], List[List[str]]]:
    lines = [ln.rstrip() for ln in md.splitlines() if ln.strip()]
    # Find header line starting with '|'
    header_idx = None
    for i, ln in enumerate(lines):
        if ln.startswith('|') and ln.count('|') > 2:
            header_idx = i
            break
    if header_idx is None:
        return [], []
    header_line = lines[header_idx]
    sep_line = lines[header_idx + 1] if header_idx + 1 < len(lines) else ''
    # Data lines until a blank or non '|' line
    data_lines = []
    for ln in lines[header_idx + 2:]:
        if not ln.startswith('|'):
            break
        data_lines.append(ln)
    def split_row(row: str) -> List[str]:
        parts = [c.strip() for c in row.strip('|').split('|')]
        return parts
    headers = split_row(header_line)
    # Skip separator row if it looks like the markdown dashes row
    body = [split_row(dl) for dl in data_lines]
    # Filter out rows that look like header/separator accidentally included
    body = [r for r in body if any(cell.strip() for cell in r) and r != headers]
    return headers, body

# ---------------------------------------------------------------------------

def build_docx(output_docx: str, base_headers: List[str], base_rows: List[List[str]], table_headers: List[str], table_rows: List[List[str]], image_paths: List[str]):
    if not ensure_package('docx', 'python-docx'):
        print('python-docx not available; cannot write Word document.')
        return
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore

    doc = Document()
    doc.add_heading('GPS Run Report', level=1)
    # Removed generator/log lines per requirement
    # Base summary first
    if base_headers and base_rows:
        doc.add_heading('Run Summary', level=2)
        # Expect single row
        base_map = {h: base_rows[0][i] if i < len(base_rows[0]) else '' for i,h in enumerate(base_headers)}
        # Build 3-row, 4-column table (label, value, label, value)
        summary_pairs = [
            ('LOCO_PILOT', base_map.get('LOCO_PILOT',''), 'SECTION', base_map.get('SECTION','')),
            ('TRAIN_NUMBER', base_map.get('TRAIN_NUMBER',''), 'DATE', base_map.get('DATE','')),
            ('LOCO_NUMBER', base_map.get('LOCO_NUMBER',''), 'TIMING', base_map.get('TIMING','')),
        ]
        t0 = doc.add_table(rows=0, cols=4)
        for a_label, a_val, b_label, b_val in summary_pairs:
            row = t0.add_row().cells
            row[0].text = a_label
            row[1].text = a_val
            row[2].text = b_label
            row[3].text = b_val
    elif base_headers:
        # Fallback to previous layout if somehow multiple rows
        t0 = doc.add_table(rows=1, cols=len(base_headers))
        for i,h in enumerate(base_headers):
            t0.rows[0].cells[i].text = h
        for rowdata in base_rows:
            r = t0.add_row().cells
            for i,val in enumerate(rowdata):
                if i < len(r):
                    r[i].text = val
    # Halt arrival table
    if table_headers:
        doc.add_heading('Halt Arrival Table', level=2)
        t = doc.add_table(rows=1, cols=len(table_headers))
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(table_headers):
            hdr_cells[i].text = h
        for row in table_rows:
            row_cells = t.add_row().cells
            for i, val in enumerate(row):
                if i < len(row_cells):
                    row_cells[i].text = val
    else:
        doc.add_paragraph('No halt arrival table parsed.')
    # Charts
    if image_paths:
        doc.add_heading('Charts', level=2)
        for img in image_paths:
            try:
                doc.add_picture(img, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f'Image embed error: {e}')
    else:
        doc.add_paragraph('No charts to embed.')
    # Removed 'End of report.' line
    doc.save(output_docx)
    # Make sure file is writable (not read-only) and remove macOS quarantine attribute if present
    try:
        import stat
        st_mode = os.stat(output_docx).st_mode
        # Add user write bit
        os.chmod(output_docx, st_mode | stat.S_IWUSR)
    except Exception:
        pass
    if sys.platform == 'darwin':
        try:
            subprocess.run(['xattr', '-d', 'com.apple.quarantine', output_docx], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception:
            pass
    print('Wrote report', output_docx)

# ---------------------------------------------------------------------------

def extract_pilot_name(csv_path: str) -> str:
    base = os.path.basename(csv_path)
    base_no_ext = os.path.splitext(base)[0]
    idx = base_no_ext.lower().find('primarygpsdata')
    if idx != -1:
        prefix = base_no_ext[:idx]
    else:
        prefix = base_no_ext
    # Remove common separators
    prefix = prefix.replace('_',' ').replace('-',' ')  # keep spaces
    # Extract alpha sequences
    import re
    parts = re.findall(r'[A-Za-z]+', prefix)
    if not parts:
        return 'Pilot'
    # Preserve single-letter uppercase (e.g., initial)
    cleaned = [p.upper() if len(p)==1 else p.capitalize() for p in parts]
    return ' '.join(cleaned).strip()

def main():
    import argparse
    ap = argparse.ArgumentParser(description='Run chart & halt table scripts and compile a Word report.')
    ap.add_argument('csv', help='Input GPS CSV file')
    ap.add_argument('--output', default='', help='Output Word document (default: dynamic name)')
    ap.add_argument('--loco', default='', help='Loco number (passed to passingHome.py)')
    ap.add_argument('--min-stop', type=int, default=30, help='Minimum stop duration seconds (passed to passingHome.py)')
    ap.add_argument('--keep-png', action='store_true', help='Keep intermediate PNG copies created from SVG.')
    ap.add_argument('--train', default='', help='Train number (passed to baseInfo.py)')
    args = ap.parse_args()

    csv_path = args.csv
    if not os.path.isfile(csv_path):
        print('Input CSV not found:', csv_path)
        sys.exit(1)
    # Determine output filename if not specified
    if not args.output:
        pilot = extract_pilot_name(csv_path)
        today = datetime.date.today()
        month_year = today.strftime('%B %Y')
        out_name = f'SPM Chart Analysis of {pilot} {month_year}.docx'
    else:
        out_name = args.output

    # Track files before running chart
    before_svgs = set(glob.glob('*.svg'))
    before_pngs = set(glob.glob('*.png'))

    # Run chart.py
    if not os.path.isfile('chart.py'):
        print('chart.py not found in current directory.')
    else:
        run_script_collect_output('chart.py', [csv_path])

    # Gather new SVG and PNG files (PNG generated directly by chart.py)
    after_svgs = set(glob.glob('*.svg'))
    after_pngs = set(glob.glob('*.png'))
    new_svgs = sorted(after_svgs - before_svgs)
    new_pngs = sorted(after_pngs - before_pngs)

    # Run passingHome.py and capture output
    table_headers: List[str] = []
    table_rows: List[List[str]] = []
    if not os.path.isfile('passingHome.py'):
        print('passingHome.py not found in current directory.')
    else:
        rc, out = run_script_collect_output('passingHome.py', [csv_path, '--loco', args.loco, '--min-stop', str(args.min_stop)])
        table_headers, table_rows = parse_markdown_table(out)
        if not table_headers:
            print('Could not parse table from passingHome.py output.')

    # Use PNGs directly (no cairosvg conversion to avoid native cairo dependency)
    pngs: List[str] = new_pngs

    # Run baseInfo.py FIRST
    base_headers: List[str] = []
    base_rows: List[List[str]] = []
    if not os.path.isfile('baseInfo.py'):
        print('baseInfo.py not found in current directory.')
    else:
        base_args = [csv_path]
        if args.loco:
            base_args += ['--loco', args.loco]
        if args.train:
            base_args += ['--train', args.train]
        rc_b, out_b = run_script_collect_output('baseInfo.py', base_args)
        bh, br = parse_markdown_table(out_b)
        base_headers, base_rows = bh, br
        if not base_headers:
            print('Could not parse base summary table from baseInfo.py output.')

    build_docx(out_name, base_headers, base_rows, table_headers, table_rows, pngs)

    # Delete generated chart image files
    for f in new_svgs + new_pngs:
        try:
            os.remove(f)
        except Exception:
            pass

    if not args.keep_png:
        # No temp dir used for conversion now; safeguard if earlier leftover
        pass

if __name__ == '__main__':
    main()
